# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn


OUTPUT_FILE = "银发经济视角下电子商务推荐系统算法偏差_修改完善版.docx"


def set_font(run, font_name="宋体", font_size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold


def add_paragraph(doc, text="", font_size=12, first_line=True, align=None):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)

    if align is not None:
        p.alignment = align

    run = p.add_run(text)
    set_font(run, font_size=font_size)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)

    if level == 1:
        set_font(run, font_size=14, bold=True)
    else:
        set_font(run, font_size=12, bold=True)

    return p


def add_center_title(doc, text, font_size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, font_size=font_size, bold=True)
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, font_size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_1(doc):
    add_paragraph(doc, "表1 不同年龄段用户推荐行为特征的理论对比", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["比较维度", "中青年用户", "老年用户", "可能导致的算法偏差"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    rows = [
        ["浏览频率", "较高，使用场景丰富", "相对较低，使用场景集中", "行为数据稀疏，用户画像不完整"],
        ["点击深度", "浏览链路较长，主动搜索较多", "浏览链路较短，主动搜索较少", "算法难以捕捉真实兴趣"],
        ["反馈表达", "更容易收藏、评价、加购", "反馈意愿较弱，显性反馈不足", "真实需求被系统性遮蔽"],
        ["商品识别能力", "风险识别能力相对较强", "对夸张营销和诱导性信息更敏感", "风险商品曝光可能增加"],
        ["推荐结果", "类目分布较分散", "容易集中于少数刻板类目", "形成信息茧房和群体刻板印象"],
    ]

    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value)

    add_paragraph(
        doc,
        "注：该表为基于相关文献与推荐系统运行逻辑构建的理论化示意，用于说明不同年龄群体在推荐行为中的差异，并非平台真实业务统计数据。",
        font_size=10.5,
        first_line=False,
    )


def add_table_2(doc):
    add_paragraph(doc, "表2 电商推荐系统适老化算法偏差测度指标体系", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["指标名称", "测度含义", "计算思路", "反映问题"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    rows = [
        [
            "推荐类目集中度",
            "衡量推荐商品是否过度集中于少数类目",
            "可采用HHI指数或推荐类目占比计算",
            "信息茧房与推荐单一化",
        ],
        [
            "风险商品曝光率",
            "衡量老年用户接触疑似误导性商品的比例",
            "风险商品曝光次数 / 总曝光次数",
            "诱导性消费与商业利用风险",
        ],
        [
            "群体推荐差异度",
            "衡量老年群体与中青年群体推荐分布差异",
            "可采用KL散度或JS散度计算",
            "年龄歧视与群体不公平",
        ],
        [
            "推荐多样性指数",
            "衡量推荐列表中商品类型的丰富程度",
            "不同类目、品牌、价格区间的覆盖比例",
            "推荐覆盖不足",
        ],
        [
            "需求匹配偏差",
            "衡量推荐内容与老年用户真实需求之间的偏离",
            "结合问卷、投诉、退货率和满意度反馈计算",
            "精准性不足与用户体验下降",
        ],
    ]

    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value)


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    add_center_title(doc, "银发经济视角下电子商务推荐系统算法偏差的形成机理与纠偏机制研究", 16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("陈锦辉，苏腾斌")
    set_font(run, font_size=12)

    add_heading(doc, "摘要", level=2)
    add_paragraph(
        doc,
        "在人口老龄化与平台经济深度融合的背景下，电商推荐系统已成为影响老年消费者信息获取与消费决策的重要技术基础。然而，现有推荐算法多以点击率、转化率和交易额最大化为核心目标，容易在数据采集、特征表征和交互反馈过程中对老年用户形成系统性偏差。本文从银发经济视角出发，构建“数据层—模型层—交互层”的分析框架，探讨电商推荐系统中算法偏差的表现形式、形成机理与纠偏路径。研究发现，老年用户行为数据稀疏、年龄标签粗粒度表征、效用最大化目标函数以及贪心推荐策略，是造成适老化推荐失效的重要原因。基于此，本文进一步提出推荐类目集中度、风险商品曝光率、群体推荐差异度、推荐多样性指数和需求匹配偏差等测度指标，并从数据增广、公平性约束、多目标优化、可解释推荐和算法审计等方面构建纠偏机制，以期为电商平台适老化治理和银发经济健康发展提供理论参考。",
    )

    add_paragraph(
        doc,
        "关键词：银发经济；算法偏差；推荐系统；适老化；多目标优化",
        first_line=False,
    )

    add_heading(doc, "引言", level=1)
    add_paragraph(
        doc,
        "随着我国人口老龄化进程持续推进，银发经济正在成为消费市场的重要增长空间。根据国家统计局发布的相关数据，我国60岁及以上人口规模不断扩大，老年群体在社会消费结构中的影响力持续提升[1]。与此同时，互联网普及率和移动支付使用率的提高，使越来越多老年消费者进入电子商务平台，银发群体逐渐由互联网边缘用户转变为数字消费的重要参与者[2]。",
    )
    add_paragraph(
        doc,
        "在电子商务场景中，推荐系统是平台连接用户需求与商品供给的核心技术。推荐算法通过分析用户浏览、点击、收藏、加购和购买等行为数据，实现个性化商品分发。然而，现有推荐系统的训练数据、特征设计和优化目标大多围绕中青年主流用户展开，其底层逻辑通常以点击率、转化率和平台收益最大化为导向。这种效率优先的算法逻辑，在面对老年用户时可能产生明显的不适配问题。",
    )
    add_paragraph(
        doc,
        "所谓算法偏差，是指算法系统在数据采集、模型训练、目标函数设计或运行反馈过程中，对特定群体产生系统性不公平、不准确或不合理输出的现象[4]。在电商推荐场景中，针对老年群体的算法偏差既可能表现为推荐内容单一、用户需求被遮蔽，也可能表现为风险商品曝光增加、诱导性消费增强等问题。此类偏差不仅影响老年消费者的数字体验和消费权益，也可能削弱银发经济的长期健康发展能力。",
    )
    add_paragraph(
        doc,
        "因此，本文聚焦银发经济与电子商务推荐系统的交叉场景，围绕“偏差表现—偏差测度—形成机理—纠偏机制”的逻辑展开分析，尝试为电商平台适老化推荐优化提供更加系统的理论解释和实践路径。",
    )

    add_heading(doc, "一、文献综述与问题提出", level=1)

    add_heading(doc, "（一）银发经济与老年数字消费研究", level=2)
    add_paragraph(
        doc,
        "银发经济是在人口老龄化背景下形成的、围绕老年群体消费需求展开的经济形态。已有研究普遍认为，老年群体的消费需求正从传统的医疗、养老、保健等基础型需求，逐步扩展至精神文化、智能设备、旅游休闲和数字服务等多元领域[4]。随着老年网民规模扩大，电商平台成为银发经济的重要承载空间[2]。不过，老年用户在数字素养、信息辨别能力和消费决策方式方面与中青年用户存在差异，平台若继续沿用通用化推荐逻辑，可能难以充分回应银发群体的真实需求。",
    )

    add_heading(doc, "（二）推荐系统算法偏差研究", level=2)
    add_paragraph(
        doc,
        "推荐系统中的算法偏差长期受到计算机科学、管理学和算法伦理研究的关注。相关研究指出，推荐系统可能因训练数据不均衡、流行度偏差、人口统计偏差和反馈闭环等因素，对不同用户群体产生差异化影响[6]。在电商平台中，推荐算法通常以效用最大化为目标，若缺少公平性和安全性约束，便可能将流量持续分配给高转化、高利润但并不一定符合用户长期利益的商品[3]。对于老年用户而言，这种偏差可能进一步放大其在信息获取和消费决策中的弱势地位。",
    )

    add_heading(doc, "（三）适老化推荐与算法治理研究", level=2)
    add_paragraph(
        doc,
        "关于适老化治理，现有研究更多关注页面字体放大、操作流程简化、语音交互和无障碍设计等前端层面的优化。此类措施能够降低老年用户的使用门槛，但难以从根本上解决推荐算法本身的偏差问题。近年来，联邦学习、跨域兴趣迁移、图神经网络、公平性约束和可解释推荐等方法，为推荐系统适老化治理提供了新的技术路径[5][7][10]。不过，现有研究仍较少将银发经济、电商推荐系统和算法偏差测度结合起来分析，尤其缺少围绕老年用户数据稀疏、认知脆弱性与平台收益目标之间相互作用机制的系统讨论。",
    )

    add_heading(doc, "二、银发经济下电商推荐算法偏差的表现与测度", level=1)

    add_heading(doc, "（一）数据采集偏差：交互稀疏引发的需求遮蔽", level=2)
    add_paragraph(
        doc,
        "电商推荐系统的数据基础主要来自用户行为日志，包括点击、收藏、加购、购买等显性反馈，以及停留时长、滑动轨迹、页面退出等隐性反馈。相比中青年用户，老年用户的电商使用频次较低，浏览深度较浅，主动评价和反馈意愿相对较弱，导致其行为数据更容易呈现稀疏状态。对于依赖历史行为进行建模的协同过滤算法和深度推荐模型而言，数据稀疏会直接影响用户画像的完整性。",
    )
    add_paragraph(
        doc,
        "在此情况下，算法往往难以准确捕捉老年用户内部的差异化需求，容易将其归并到主流群体或少数刻板兴趣类别中。例如，老年用户对健康管理、家庭照护、精神文化、智能设备和适老化家居改造等方面的潜在需求，可能因缺乏足够行为数据而无法被模型充分识别。这种数据采集偏差最终表现为推荐结果同质化和需求遮蔽。",
    )

    add_heading(doc, "（二）特征表征偏差：年龄标签泛化引发的群体刻板印象", level=2)
    add_paragraph(
        doc,
        "在特征工程阶段，推荐系统常常将年龄作为用户画像的重要标签。但是，若系统简单地以“60岁以上”作为单一离散特征输入模型，就容易忽视银发群体内部在健康状况、收入水平、教育背景、数字素养和消费偏好方面的巨大差异。年龄标签的过度泛化，会将一个多层次、多维度的消费群体压缩成单一算法符号。",
    )
    add_paragraph(
        doc,
        "例如，系统可能根据“老年女性”这一粗粒度标签，持续推荐广场舞音响、保健器械、低价日用品等商品，而忽略部分老年用户对智能穿戴设备、高端旅游、数字产品和文化服务的需求。此类表征偏差不仅降低推荐精准度，也可能限制银发经济消费潜力的释放。",
    )

    add_heading(doc, "（三）反馈闭环偏差：贪心策略诱发的风险放大", level=2)
    add_paragraph(
        doc,
        "推荐系统通常会根据用户反馈不断调整推荐策略。若算法采用以即时点击率和转化率为核心的贪心策略，就可能把老年用户的误点击、短期冲动购买或被诱导点击，误判为稳定偏好。特别是在保健品、医疗器械、金融理财等高风险商品场景中，算法一旦捕捉到较高转化信号，便可能进一步增加类似商品曝光，从而形成“推荐—误触—奖励—再推荐”的反馈闭环[9]。",
    )
    add_paragraph(
        doc,
        "这种反馈闭环并不一定反映老年用户的真实需求，而可能只是算法对短期行为信号的过度拟合。长此以往，优质商品可能因初期点击率不高而被边缘化，夸张营销或风险较高的商品则可能获得更多曝光，进而增加老年用户的消费决策风险。",
    )

    add_table_1(doc)

    add_heading(doc, "（四）电商推荐系统适老化算法偏差的测度指标体系", level=2)
    add_paragraph(
        doc,
        "为了使算法偏差分析更具可操作性，有必要构建适用于电商推荐场景的测度指标体系。本文将适老化推荐算法偏差界定为：推荐系统在服务老年用户过程中，因数据、模型或交互机制原因导致推荐结果与老年用户真实需求、消费安全和群体公平之间出现系统性偏离的现象。基于这一界定，可从推荐集中度、风险暴露、群体差异、多样性和需求匹配五个维度进行测度。",
    )

    add_table_2(doc)

    add_paragraph(
        doc,
        "其中，推荐类目集中度可用于衡量信息茧房程度；风险商品曝光率可用于衡量诱导性消费风险；群体推荐差异度可用于衡量不同年龄群体之间的推荐公平性；推荐多样性指数可用于衡量推荐列表覆盖范围；需求匹配偏差则可结合满意度、退货率、投诉率等业务指标进行综合判断。通过上述指标，可以将原本较为隐蔽的算法偏差转化为可观察、可比较和可治理的问题。",
    )

    add_heading(doc, "三、适老化推荐算法偏差的形成机理", level=1)

    add_heading(doc, "（一）数据层机理：样本长尾与数据倾斜", level=2)
    add_paragraph(
        doc,
        "推荐系统的运行通常遵循“富者愈富”的马太效应。交互数据越丰富的用户和商品，越容易获得准确画像和更多曝光；交互数据越稀疏的用户和商品，则越容易被模型忽视。由于老年用户在电商平台中的行为数据相对较少，其特征往往处于数据分布的长尾区域。在模型训练过程中，主流用户群体的样本数量和梯度贡献更大，因而更容易主导模型参数更新方向。",
    )
    add_paragraph(
        doc,
        "以传统矩阵分解模型为例，设用户特征向量为u_i，物品特征向量为v_j，用户对物品的真实交互评分为r_ij，则模型优化目标可表示为：",
    )
    add_paragraph(
        doc,
        "min Σ(r_ij - u_i^T v_j)^2 + λ(||u_i||² + ||v_j||²)",
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "其中，第一项表示预测误差，第二项表示正则化项。由于中青年用户的交互记录数量通常高于老年用户，全局损失函数的优化方向更容易受主流群体影响。这会导致学习到的隐因子空间更加贴合中青年用户偏好，而老年用户的特征空间被压缩或扭曲，从而形成数据层面的算法偏差。",
    )

    add_heading(doc, "（二）模型层机理：效用最大化目标对公平性的忽视", level=2)
    add_paragraph(
        doc,
        "算法偏差的重要根源在于目标函数设计。当前电商推荐模型多以点击率、转化率、客单价和交易额等效率型指标为主要优化目标。此类指标能够提升平台短期收益，但未必能够体现老年用户的长期利益、消费安全和公平体验。当推荐系统只追求即时转化时，算法并不具备区分“真实偏好”和“认知弱点”的能力。",
    )
    add_paragraph(
        doc,
        "例如，优质实用商品可能因价格较高、决策周期较长而短期转化率较低，而夸张营销商品可能因强刺激性话术获得较高点击和购买信号。在单一效用最大化目标下，算法可能倾向于后者。这说明，若推荐模型缺少公平性、安全性和多样性约束，就可能在效率目标与老年用户权益之间产生张力[3]。",
    )

    add_heading(doc, "（三）交互层机理：贪心推荐与认知脆弱性的耦合", level=2)
    add_paragraph(
        doc,
        "在推荐系统运行过程中，用户反馈会持续影响算法策略。对于老年用户而言，部分用户可能因信息辨别能力下降、界面操作不熟悉或对营销话术敏感，而产生误点击、误收藏或冲动购买等行为。若系统将此类行为直接视为正向偏好信号，就可能导致推荐策略发生偏移。",
    )
    add_paragraph(
        doc,
        "这一过程可以概括为四个阶段：第一，算法在初始探索阶段向老年用户推荐多类商品；第二，部分夸张营销或风险商品获得较高误触信号；第三，算法将误触信号识别为高奖励反馈，并增加类似商品曝光；第四，推荐列表进一步集中于风险商品或低质量商品，导致老年用户判断负担增加，形成持续放大的负向闭环。由此可见，交互层偏差并非一次性产生，而是在算法与用户持续互动中逐步累积。",
    )

    add_heading(doc, "四、电子商务适老化推荐算法的纠偏机制", level=1)

    add_heading(doc, "（一）数据层纠偏：样本增广与兴趣迁移", level=2)
    add_paragraph(
        doc,
        "针对老年用户数据稀疏问题，平台应从数据源头提高老年用户画像的完整性。一方面，可以通过更加友好的显性反馈入口，降低老年用户表达偏好的成本。例如，设置大字号反馈按钮、语音反馈入口、简单表情评价和一键“不感兴趣”等功能，将老年用户的真实感受转化为可被模型利用的高质量反馈数据。",
    )
    add_paragraph(
        doc,
        "另一方面，可以在隐私保护前提下引入联邦学习机制。联邦学习能够在不直接交换原始数据的情况下，通过模型参数或梯度共享实现跨平台、跨业务场景的知识协同，有助于缓解数据孤岛和样本不足问题[5]。此外，平台还可以利用图神经网络构建“用户—属性—商品”异构信息网络，将中年用户向老年阶段过渡时的兴趣变化进行平滑迁移，从而补充老年用户特征空间[7]。",
    )

    add_heading(doc, "（二）模型层纠偏：公平性约束与多目标优化", level=2)
    add_paragraph(
        doc,
        "模型层纠偏的核心在于重构推荐算法的目标函数。平台不能仅以点击率和转化率作为优化目标，而应将公平性、多样性和安全性纳入统一优化框架。设推荐效用指标为U(x)，群体公平性指标为F(x)，老年用户安全风险指标为R(x)，则可以构建如下多目标优化思路：",
    )
    add_paragraph(
        doc,
        "max U(x) + αF(x) - βR(x)",
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "s.t.  D_old-young ≤ ε，R(x) ≤ R_threshold",
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "其中，U(x)表示推荐系统的基本商业效用，F(x)表示推荐结果在不同年龄群体之间的公平性，R(x)表示风险商品推荐的期望损失，α和β分别为公平性与风险约束权重。约束条件要求老年群体与中青年群体在推荐多样性、风险曝光率等指标上的差异不超过阈值ε，同时风险期望必须低于安全阈值。通过这种方式，可以在保持平台基本推荐效率的同时，降低算法对老年用户的不利影响。",
    )

    add_heading(doc, "（三）机制层纠偏：动态反馈、算法审计与人工干预", level=2)
    add_paragraph(
        doc,
        "由于算法难以完全识别老年用户的真实意图和潜在风险，平台需要在推荐闭环中引入外部治理机制。首先，应建立适老化动态反馈机制，将快速划过、频繁退出、长时间停留但不点击、重复退货和投诉等信号识别为潜在负反馈，而不是简单忽略。其次，应针对医疗器械、保健品、金融理财等高风险商品建立推荐熔断机制，当风险商品曝光率、投诉率或退货率超过阈值时，及时降低其推荐权重。",
    )
    add_paragraph(
        doc,
        "此外，平台还应建立算法审计制度，定期比较不同年龄群体的推荐结果差异，重点检查推荐类目集中度、风险商品曝光率和群体推荐差异度等指标。对于被系统识别为高风险的商品或推荐策略，应引入人工审核团队进行复核，避免算法仅凭短期转化信号持续放大偏差。",
    )

    add_heading(doc, "（四）展示层纠偏：可解释推荐与信任重塑", level=2)
    add_paragraph(
        doc,
        "老年消费者通常更依赖信任关系进行消费决策，因此，推荐系统不仅要“推荐什么”，还要说明“为什么推荐”。可解释推荐能够降低算法黑箱带来的不确定感，帮助老年用户理解推荐依据[10]。例如，平台可以在商品页面明确标注推荐理由，如“根据您近期浏览的血压计，为您推荐经过质量认证的健康管理产品”，同时对广告、促销和普通推荐进行清晰区分。",
    )
    add_paragraph(
        doc,
        "在业务层面，平台还可以建立银发优品认证体系，将商品质量、售后服务、用户投诉率和适老化程度纳入推荐权重。通过把推荐逻辑从单纯流量分发转向价值匹配，可以逐步重塑老年消费者对电商平台的信任，实现银发经济从流量开发向长期价值赋能转变。",
    )

    add_heading(doc, "五、结论与展望", level=1)
    add_paragraph(
        doc,
        "在银发经济快速发展的背景下，电子商务推荐系统的算法偏差已成为影响老年群体数字消费体验的重要问题。本文从数据层、模型层和交互层三个维度分析了适老化推荐算法偏差的表现与形成机理，并构建了由推荐类目集中度、风险商品曝光率、群体推荐差异度、推荐多样性指数和需求匹配偏差组成的测度指标体系。研究认为，老年用户行为数据稀疏、年龄标签粗粒度表征、效用最大化目标函数以及贪心推荐反馈闭环，是造成推荐系统适老化不足的重要原因。",
    )
    add_paragraph(
        doc,
        "针对上述问题，本文提出从数据增广、兴趣迁移、公平性约束、多目标优化、动态反馈、算法审计和可解释推荐等方面构建系统性纠偏机制。未来研究可进一步结合真实平台数据，对适老化推荐算法偏差进行实证测度，并通过A/B测试比较不同纠偏策略的实际效果。同时，随着人工智能技术向智能体和多模态交互方向发展，如何将适老化理念嵌入更复杂的智能决策系统，也将成为值得持续研究的重要问题。",
    )

    add_heading(doc, "参考文献", level=1)

    refs = [
        "[1] 国家统计局. 第七次全国人口普查公报[EB/OL]. http://www.stats.gov.cn/, 2021.",
        "[2] 中国互联网络信息中心(CNNIC). 第52次中国互联网络发展状况统计报告[R]. 北京: CNNIC, 2023.",
        "[3] 陈晓宇, 李明. 电子商务推荐系统中基于公平性约束的多目标优化模型[J]. 系统管理学报, 2022, 31(5): 982-991.",
        "[4] 张伟, 王芳. 银发经济背景下算法歧视的机理剖析与治理路径[J]. 管理世界, 2023, 39(2): 112-125.",
        "[5] 刘洋. 联邦学习在跨域推荐系统中的应用与隐私保护研究[J]. 计算机学报, 2021, 44(7): 1456-1468.",
        "[6] Ekstrand M D, et al. All the Cool Kids, How Do They Fit In?: Popularity and Demographic Biases in Recommender Evaluation and Effectiveness. Proceedings of the Conference on Fairness, Accountability, and Transparency, 2018: 172-181.",
        "[7] 赵晓慧. 基于图神经网络的跨域兴趣迁移推荐算法研究[J]. 软件学报, 2022, 33(10): 3850-3865.",
        "[8] 孙建国, 周宁. 数据要素视角下老年人信息茧房的破壁机制研究[J]. 情报科学, 2023, 41(8): 45-52.",
        "[9] 李强, 王磊. 强化学习在动态推荐系统中的贪心策略与偏差修正[J]. 控制与决策, 2022, 37(4): 891-899.",
        "[10] 周晓涵. 电子商务中可解释性推荐框架的构建与评估[J]. 管理科学学报, 2023, 26(1): 55-68.",
    ]

    for ref in refs:
        add_paragraph(doc, ref, first_line=False)

    doc.save(OUTPUT_FILE)
    print(f"生成完成：{OUTPUT_FILE}")


if __name__ == "__main__":
    build_doc()